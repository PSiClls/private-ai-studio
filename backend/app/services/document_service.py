import os
import uuid
import asyncio
from typing import List, Optional
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ..models import Document, DocumentChunk
from .embeddings import embedding_service, vector_store


class TextSplitter:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            if end >= len(text):
                chunks.append(text[start:])
                break

            chunk = text[start:end]
            last_period = chunk.rfind(".")
            last_newline = chunk.rfind("\n")
            split_at = max(last_period, last_newline)
            if split_at > self.chunk_size // 2:
                end = start + split_at + 1
                chunks.append(text[start:end])
            else:
                chunks.append(chunk)

            start = end - self.chunk_overlap
        return [c.strip() for c in chunks if c.strip()]


class DocumentProcessor:
    async def process_file(
        self,
        file_path: str,
        filename: str,
        db: AsyncSession,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> Document:
        ext = os.path.splitext(filename)[1].lower()
        text = await self._extract_text(file_path, ext)

        splitter = TextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = splitter.split_text(text)

        doc = Document(
            filename=filename,
            file_type=ext,
            file_size=os.path.getsize(file_path),
            chunk_count=len(chunks),
            status="processing",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        try:
            embeddings = await embedding_service.embed(chunks)

            chunk_ids = []
            metadatas = []
            db_chunks = []

            for i, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"chunk_{doc.id}_{i}_{uuid.uuid4().hex[:8]}"
                chunk_ids.append(chunk_id)

                metadata = {
                    "chunk_index": i,
                    "filename": filename,
                    "page_number": self._estimate_page(i, chunk_size),
                }
                metadatas.append(metadata)

                db_chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=i,
                    text_preview=chunk_text[:200],
                    token_count=len(chunk_text.split()),
                    chroma_id=chunk_id,
                )
                db_chunks.append(db_chunk)

            await vector_store.add_chunks(
                doc_id=doc.id,
                chunk_ids=chunk_ids,
                embeddings=embeddings,
                metadatas=metadatas,
                texts=chunks,
            )

            for db_chunk in db_chunks:
                db.add(db_chunk)

            doc.status = "ready"
            await db.commit()
            await db.refresh(doc)

        except Exception as e:
            doc.status = "error"
            await db.commit()
            raise e

        return doc

    def _estimate_page(self, chunk_index: int, chunk_size: int) -> int:
        tokens_per_page = 500
        return (chunk_index * chunk_size) // tokens_per_page + 1

    async def _extract_text(self, file_path: str, ext: str) -> str:
        if ext in (".txt", ".md"):
            return await self._read_text(file_path)
        elif ext == ".csv":
            return await self._read_csv(file_path)
        elif ext == ".pdf":
            return await self._read_pdf(file_path)
        elif ext == ".docx":
            return await self._read_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def _read_text(self, file_path: str) -> str:
        loop = asyncio.get_event_loop()
        def _read():
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        return await loop.run_in_executor(None, _read)

    async def _read_csv(self, file_path: str) -> str:
        import csv
        loop = asyncio.get_event_loop()
        def _read():
            rows = []
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(", ".join(row))
            return "\n".join(rows)
        return await loop.run_in_executor(None, _read)

    async def _read_pdf(self, file_path: str) -> str:
        loop = asyncio.get_event_loop()
        def _read():
            try:
                import fitz
                text = ""
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text() + "\n"
                return text
            except ImportError:
                pass
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                pass
            raise ValueError("No PDF reader available (install pymupdf or pdfplumber)")
        return await loop.run_in_executor(None, _read)

    async def _read_docx(self, file_path: str) -> str:
        loop = asyncio.get_event_loop()
        def _read():
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        return await loop.run_in_executor(None, _read)

    async def delete_document(self, doc: Document, db: AsyncSession):
        await vector_store.delete_collection(doc.id)
        result = await db.execute(
            select(DocumentChunk).where(DocumentChunk.document_id == doc.id)
        )
        chunks = result.scalars().all()
        for chunk in chunks:
            await db.delete(chunk)
        await db.delete(doc)
        await db.commit()


document_processor = DocumentProcessor()
